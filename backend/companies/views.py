from rest_framework import generics, viewsets, status
from rest_framework.response import Response
from django.http import HttpResponse
from rest_framework.views import APIView
from .models import Company, Inspection
from .serializers import CompanySerializer, InspectionSerializer
from docx import Document  # Ensure you have python-docx installed
from io import BytesIO

class CompanyListCreate(generics.ListCreateAPIView):
    queryset = Company.objects.all()
    serializer_class = CompanySerializer

class CompanyDetail(generics.RetrieveAPIView):
    queryset = Company.objects.all()
    serializer_class = CompanySerializer

class InspectionViewSet(viewsets.ModelViewSet):
    queryset = Inspection.objects.all()
    serializer_class = InspectionSerializer

    def create(self, request, *args, **kwargs):
        aktenzeichen = request.data.get('aktenzeichen')
        try:
            company = Company.objects.get(aktenzeichen=aktenzeichen)
            inspection = Inspection.objects.create(company=company, **request.data)
            return Response({'status': 'Inspection created'})
        except Company.DoesNotExist:
            return Response({'error': 'Company with the given Aktenzeichen does not exist'}, status=400)

class GenerateInspectionDoc(APIView):
    def post(self, request, *args, **kwargs):
        # Get the data from the request
        inspection_data = request.data
        
        # Validate and serialize the data
        serializer = InspectionSerializer(data=inspection_data)
        if serializer.is_valid():
            # Save the inspection data to the backend
            serializer.save()

            # Create a new Document
            doc = Document()
            doc.add_heading('Inspection Report', level=1)

            # Add inspection data to the document
            doc.add_paragraph(f"Company: {inspection_data.get('company', 'N/A')}")
            doc.add_paragraph(f"Inspection Date: {inspection_data.get('inspection_date', 'N/A')}")
            doc.add_paragraph(f"Inspector Name: {inspection_data.get('inspector_name', 'N/A')}")
            doc.add_paragraph(f"Team Lead: {inspection_data.get('team_lead', 'N/A')}")
            doc.add_paragraph(f"Additional Inspectors: {inspection_data.get('additional_inspectors', 'N/A')}")
            doc.add_paragraph(f"Authority: {inspection_data.get('authority', 'N/A')}")
            doc.add_paragraph(f"Reference Number: {inspection_data.get('reference_number', 'N/A')}")
            doc.add_paragraph(f"Introduction: {inspection_data.get('introduction', 'N/A')}")
            doc.add_paragraph(f"SMF Info: {inspection_data.get('smf_info', 'N/A')}")
            doc.add_paragraph(f"Previous Inspection: {inspection_data.get('previous_inspection', 'N/A')}")
            doc.add_paragraph(f"Findings: {inspection_data.get('findings', 'N/A')}")
            doc.add_paragraph(f"Active Substances: {inspection_data.get('active_substances', 'N/A')}")
            doc.add_paragraph(f"Conclusions: {inspection_data.get('conclusions', 'N/A')}")
            doc.add_paragraph(f"Quality Management: {inspection_data.get('quality_management', 'N/A')}")
            doc.add_paragraph(f"Personnel: {inspection_data.get('personnel', 'N/A')}")
            doc.add_paragraph(f"Equipment: {inspection_data.get('equipment', 'N/A')}")
            doc.add_paragraph(f"Documentation: {inspection_data.get('documentation', 'N/A')}")
            doc.add_paragraph(f"Production: {inspection_data.get('production', 'N/A')}")
            doc.add_paragraph(f"Quality Control: {inspection_data.get('quality_control', 'N/A')}")
            doc.add_paragraph(f"Contract Testing: {inspection_data.get('contract_testing', 'N/A')}")
            doc.add_paragraph(f"Complaints: {inspection_data.get('complaints', 'N/A')}")
            doc.add_paragraph(f"Self Inspection: {inspection_data.get('self_inspection', 'N/A')}")
            doc.add_paragraph(f"Storage: {inspection_data.get('storage', 'N/A')}")
            doc.add_paragraph(f"Other Aspects: {inspection_data.get('other_aspects', 'N/A')}")
            doc.add_paragraph(f"Site Description: {inspection_data.get('site_description', 'N/A')}")
            doc.add_paragraph(f"Samples Taken: {inspection_data.get('samples_taken', 'N/A')}")
            doc.add_paragraph(f"Critical Errors: {inspection_data.get('critical_errors', 'N/A')}")
            doc.add_paragraph(f"Serious Errors: {inspection_data.get('serious_errors', 'N/A')}")
            doc.add_paragraph(f"Other Errors: {inspection_data.get('other_errors', 'N/A')}")
            doc.add_paragraph(f"Remarks: {inspection_data.get('remarks', 'N/A')}")

            # Save the document to a BytesIO object
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Create a response with the DOCX file
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename="inspection_report.docx"'
            return response
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
